"""Protokoll-Render-Service.

Reine Funktion: dict (NachtragDetailResponse-Format) -> python-docx Document.
Kein DB-Zugriff. Wird aus dem Sandbox-Skript
`C:\\Tools\\claude-sandbox\\bp-protokoll\\test_protokoll_render.py` uebernommen
und dort bei Bedarf weiterentwickelt.

G1-konform: keine Wertungen, keine Schuldzuweisungen. Inhalte sind direkt
aus dem Datensatz und werden nicht interpretiert. Begruendungen aus
nachtragspruefung werden faktentreu uebernommen.
"""

from __future__ import annotations

from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_zelle(zelle, text: str, *, bold: bool = False, size: int = 10,
               farbe: RGBColor | None = None, align=None) -> None:
    """Fuegt einen Absatz mit gewuenschtem Style in eine Tabellenzelle ein."""
    p = zelle.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.text = ""
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if farbe:
        run.font.color.rgb = farbe


def _add_paragraph(doc: Document, text: str, *, bold: bool = False,
                   size: int = 11, farbe: RGBColor | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if farbe:
        run.font.color.rgb = farbe
    return p


def render_nachtrag_protokoll(nt: dict[str, Any]) -> Document:
    """Erzeugt das Word-Protokoll fuer einen Nachtrag (AP 2.5).

    Args:
        nt: Dict im Format der NachtragDetailResponse (aus
            NachtragsService.lade_nachtrag). Wird nicht mutiert.

    Returns:
        Document — kann ueber doc.save(buf) in einen BytesIO geschrieben
        und vom FastAPI-Router als StreamingResponse zurueckgegeben werden.
    """
    doc = Document()

    # ----- Seitenraender -----
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ----- Titel + Untertitel -----
    titel = doc.add_heading(f"Nachtragspruefprotokoll {nt.get('nummer', '')}", level=0)
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(nt.get("gegenstand", ""))
    run.italic = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ----- 1. Stammdaten -----
    doc.add_heading("1. Stammdaten", level=2)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for label, wert in [
        ("Nachtragsnummer", nt.get("nummer", "—")),
        ("Gegenstand", nt.get("gegenstand", "—")),
        ("Status", nt.get("status", "—")),
        ("Variante", nt.get("nachtragsvariante") or "—"),
        ("Kostengruppe DIN 276", nt.get("kostengruppe_din276") or "—"),
        ("Erstellt am / von", f"{nt.get('erstellt_am', '')} / {nt.get('erstellt_von', '')}".strip(" /")),
    ]:
        r = t.add_row().cells
        _add_zelle(r[0], label, bold=True, size=10)
        _add_zelle(r[1], str(wert) if wert is not None else "—", size=10)

    doc.add_paragraph()

    # ----- 2. Dreiklang Q/Z/K (G3) -----
    doc.add_heading("2. Dreiklang Qualitaet / Zeit / Kosten", level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hd = t.rows[0].cells
    _add_zelle(hd[0], "Qualitaet (Q)", bold=True, size=10)
    _add_zelle(hd[1], "Zeit (Z) — Arbeitstage", bold=True, size=10)
    _add_zelle(hd[2], "Kosten (K) — EUR", bold=True, size=10)
    werte = t.add_row().cells
    _add_zelle(werte[0], nt.get("qualitaetsauswirkung") or "—", size=10)

    # Zeit: Vorzeichen darstellen
    z = nt.get("zeitauswirkung_tage")
    if z is None:
        z_text = "—"
    else:
        z_text = f"+{z}" if z > 0 else str(z)
    _add_zelle(werte[1], z_text, size=10)

    # Kosten: gefordert + genehmigt (deutsche Formatierung mit Tausendertrennzeichen)
    def _eur(v) -> str:
        if v is None:
            return "—"
        return f"{float(v):,.2f} EUR".replace(",", "X").replace(".", ",").replace("X", ".")

    k_text_parts = []
    if nt.get("betrag_gefordert") is not None:
        k_text_parts.append(f"gefordert: {_eur(nt['betrag_gefordert'])}")
    if nt.get("betrag_genehmigt") is not None:
        k_text_parts.append(f"genehmigt: {_eur(nt['betrag_genehmigt'])}")
    _add_zelle(werte[2], "\n".join(k_text_parts) or "—", size=10)

    doc.add_paragraph()

    # ----- 3. Pruefablauf 7 Schritte -----
    doc.add_heading("3. Pruefablauf (7-Schritte-Workflow nach AP 2.1)", level=2)
    schritte = nt.get("pruefschritte") or []
    if schritte:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        hd = t.rows[0].cells
        _add_zelle(hd[0], "Nr", bold=True, size=10)
        _add_zelle(hd[1], "Schritt", bold=True, size=10)
        _add_zelle(hd[2], "Ergebnis", bold=True, size=10)
        _add_zelle(hd[3], "Abgeschlossen", bold=True, size=10)
        for s in schritte:
            r = t.add_row().cells
            _add_zelle(r[0], str(s.get("schritt", "")), size=9)
            _add_zelle(r[1], s.get("titel") or "—", size=9)
            _add_zelle(r[2], s.get("ergebnis") or "—", size=9)
            amzeit = s.get("abgeschlossen_am")
            # ISO-String oder datetime auf TT.MM.YYYY HH:MM truncieren
            if amzeit:
                txt = str(amzeit)[:16].replace("T", " ")
            else:
                txt = "—"
            _add_zelle(r[3], txt, size=9)
    else:
        _add_paragraph(doc, "Keine Pruefschritte vorhanden.", size=10)

    doc.add_paragraph()

    # ----- 4. Entscheidung Grund / Hoehe (NT-F-04) -----
    s6 = next((s for s in schritte if s.get("schritt") == 6), None)
    if s6 and (s6.get("entscheidung_grund") is not None or s6.get("entscheidung_hoehe") is not None):
        doc.add_heading("4. Entscheidung dem Grunde / der Hoehe nach", level=2)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"

        def _add_ent_row(label: str, bool_val: bool | None, begruendung: str | None) -> None:
            r = t.add_row().cells
            _add_zelle(r[0], label, bold=True, size=10)
            if bool_val is True:
                txt = "Anerkannt"
            elif bool_val is False:
                txt = "Zurueckgewiesen / strittig"
            else:
                txt = "— nicht gesetzt"
            inhalt = txt + (f"\nBegruendung: {begruendung}" if begruendung else "")
            _add_zelle(r[1], inhalt, size=10)

        _add_ent_row("Dem Grunde nach", s6.get("entscheidung_grund"), s6.get("begruendung_grund"))
        _add_ent_row("Der Hoehe nach", s6.get("entscheidung_hoehe"), s6.get("begruendung_hoehe"))

        doc.add_paragraph()

    # ----- 5. NTV-Verknuepfung (NT-F-05) -----
    if nt.get("ntv_id"):
        doc.add_heading("5. Nachtragsvereinbarung", level=2)
        ntv_nummer = (nt.get("nummer") or "").replace("NT-", "NTV-")
        _add_paragraph(
            doc,
            f"Nach erfolgreicher Genehmigung wurde automatisch die Nachtragsvereinbarung "
            f"{ntv_nummer} im System angelegt (interne ID: {str(nt['ntv_id'])[:8]}).",
            size=10,
        )

    # ----- Footer-Vermerk -----
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')} · "
        f"BauPilot AP 2.5 · "
        f"Faktenbasiert, ohne Schuldzuweisungen (G1). Append-Only-Datenbasis (G2)."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    return doc
